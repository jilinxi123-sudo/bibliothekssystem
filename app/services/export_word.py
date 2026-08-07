from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, RGBColor

from app.models import BookOut
from app.services.export_columns import (
    column_labels,
    history_text,
    is_history_column,
    is_image_column,
    text_value,
)

COLUMN_WIDTHS_MM = {
    "isbn": 22, "title": 38, "author": 30, "publisher": 24, "published_year": 10,
    "themes": 20, "source": 20, "location": 20, "notes": 30, "due_date": 16,
    "created_at": 22, "updated_at": 24, "cover": 18, "history": 55,
}
COVER_IMG_WIDTH_MM = 15

HEADER_FILL = "6FA287"
STRIPE_FILL = "F2F8F0"


def _set_cell_background(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        if i < len(lines) - 1:
            run.add_break()


def _set_cell_image(cell, image_path: Path) -> None:
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(COVER_IMG_WIDTH_MM))


def build_word_export(
    books: list[BookOut],
    columns: list[str],
    covers_dir: Optional[Path] = None,
    history_by_isbn: Optional[dict[str, list[dict]]] = None,
) -> bytes:
    history_by_isbn = history_by_isbn or {}
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    heading = doc.add_heading("Bibliothekskatalog", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Anzahl Bücher: {len(books)}    ·    Exportiert am: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    meta_run.italic = True

    headers = column_labels(columns)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    widths = [Mm(COLUMN_WIDTHS_MM.get(key, 24)) for key in columns]

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].width = widths[i]
        _set_cell_text(header_cells[i], header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _set_cell_background(header_cells[i], HEADER_FILL)

    for index, book in enumerate(books):
        row_cells = table.add_row().cells
        for i, key in enumerate(columns):
            row_cells[i].width = widths[i]
            if is_image_column(key):
                cover_path = covers_dir / f"{book.isbn}.jpg" if covers_dir else None
                if cover_path and cover_path.exists():
                    _set_cell_image(row_cells[i], cover_path)
            elif is_history_column(key):
                _set_cell_text(row_cells[i], history_text(history_by_isbn.get(book.isbn, [])))
            else:
                _set_cell_text(row_cells[i], text_value(book, key))
        if index % 2 == 1:
            for cell in row_cells:
                _set_cell_background(cell, STRIPE_FILL)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
